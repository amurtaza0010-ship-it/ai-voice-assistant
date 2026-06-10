import logging
import os
import uuid
from typing import List, Optional

from chromadb.utils.embedding_functions import DefaultEmbeddingFunction
from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.embeddings import Embeddings
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument

from app.core.config import settings

logger = logging.getLogger(__name__)


class _ChromaLocalEmbeddings(Embeddings):
    """On-device embeddings via Chroma's default ONNX model (no API key)."""

    def __init__(self) -> None:
        self._ef = DefaultEmbeddingFunction()

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self._ef(texts)

    def embed_query(self, text: str) -> List[float]:
        return self._ef([text])[0]


def embedding_provider_name() -> str:
    if settings.openai_api_key:
        return "openai:text-embedding-3-small"
    return "chromadb:default-onnx"


def _embeddings() -> Embeddings:
    if settings.openai_api_key:
        return OpenAIEmbeddings(
            model="text-embedding-3-small",
            openai_api_key=settings.openai_api_key,
            check_embedding_ctx_length=False,
        )
    return _ChromaLocalEmbeddings()


def _collection_name(user_id: uuid.UUID) -> str:
    return f"user_{str(user_id).replace('-', '_')}"


def get_vectorstore(user_id: uuid.UUID) -> Chroma:
    os.makedirs(settings.chroma_path, exist_ok=True)
    collection = _collection_name(user_id)
    logger.info(
        "RAG vectorstore: chroma_path=%s collection=%s embedding_provider=%s",
        settings.chroma_path,
        collection,
        embedding_provider_name(),
    )
    return Chroma(
        collection_name=collection,
        embedding_function=_embeddings(),
        persist_directory=settings.chroma_path,
    )


async def ingest_file(user_id: uuid.UUID, path: str, original_name: str) -> int:
    import asyncio

    return await asyncio.to_thread(ingest_file_sync, user_id, path, original_name)


def ingest_file_sync(user_id: uuid.UUID, path: str, original_name: str) -> int:
    ext = original_name.lower().rsplit(".", 1)[-1] if "." in original_name else ""
    provider = embedding_provider_name()

    logger.info(
        "RAG ingest start: user_id=%s file=%s ext=%s path=%s embedding_provider=%s",
        user_id,
        original_name,
        ext,
        path,
        provider,
    )

    docs: List[LCDocument] = []
    if ext == "pdf":
        loader = PyPDFLoader(path)
        docs = loader.load()
    elif ext in ("txt", "md"):
        loader = TextLoader(path, encoding="utf-8")
        docs = loader.load()
    elif ext == "docx":
        docs = _load_docx(path)
    else:
        raise ValueError("Unsupported file type")

    extracted_chars = sum(len((d.page_content or "")) for d in docs)
    logger.info(
        "RAG text extraction: pages=%d extracted_text_length=%d",
        len(docs),
        extracted_chars,
    )

    if extracted_chars == 0:
        raise ValueError(
            "No text could be extracted from the document. "
            "The file may be scanned/image-only or empty."
        )

    for d in docs:
        d.metadata = d.metadata or {}
        d.metadata["source"] = original_name

    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=200)
    chunks = splitter.split_documents(docs)
    logger.info("RAG chunking: chunk_count=%d", len(chunks))

    if not chunks:
        raise ValueError("Chunking produced zero chunks")

    vs = get_vectorstore(user_id)
    vs.add_documents(chunks)
    try:
        vs.persist()
    except Exception as persist_err:
        logger.warning("RAG persist skipped: %s", persist_err)

    logger.info(
        "RAG ingest complete: user_id=%s file=%s chunks_indexed=%d vector_db=chroma collection=%s",
        user_id,
        original_name,
        len(chunks),
        _collection_name(user_id),
    )
    return len(chunks)


def _load_docx(path: str) -> List[LCDocument]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    return [LCDocument(page_content=text or "(empty document)", metadata={"source": path})]


async def retrieve_context(user_id: uuid.UUID, query: str, k: int = 6) -> Optional[str]:
    import asyncio

    try:
        vs = get_vectorstore(user_id)
        results = await asyncio.to_thread(vs.similarity_search, query, k)
        logger.info(
            "RAG retrieve: user_id=%s query_len=%d hits=%d embedding_provider=%s",
            user_id,
            len(query),
            len(results),
            embedding_provider_name(),
        )
        if not results:
            return None
        parts = []
        for r in results:
            src = r.metadata.get("source", "doc")
            parts.append(f"[{src}]\n{r.page_content}")
        return "\n\n---\n\n".join(parts)
    except Exception as e:
        logger.warning("RAG retrieve failed: %s", e, exc_info=True)
        return None


async def delete_user_collection(user_id: uuid.UUID) -> None:
    try:
        vs = get_vectorstore(user_id)
        vs.delete_collection()
    except Exception as e:
        logger.warning("delete collection: %s", e)
